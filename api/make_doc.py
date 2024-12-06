from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
from os import makedirs
from classes import Form
from logging import getLogger

logger = getLogger("uvicorn.app")
print = logger.info


def open_docx(path: str) -> DocumentObject:
    try:
        doc = Document(path)  # ファイル名とパスを正確に指定
    except FileNotFoundError:
        print("ファイルが見つかりません。ファイル名やパスを確認してください。")
        exit()
    except Exception as e:
        print(f"エラーが発生しました: {e}")
        exit()
    return doc


# Wordファイルを開く
DOC = open_docx("api/Test.docx")


def main(form: Form, doc: DocumentObject = DOC):
    # print("段落の個数:", len(doc.paragraphs))

    # print("表の個数:", len(doc.tables))

    # # Docファイルの内容を表示
    # num = 0
    # for para in doc.paragraphs:
    #     num = num + 1
    #     print(num, para.text)

    # 日付編集
    dt_now = datetime.datetime.now()
    replace(3, "yyyy", str(dt_now.year))
    replace(3, "mm", str(dt_now.month))
    replace(3, "dd", str(dt_now.day))

    # 行を挿入して文字を入力する例（５行目に新たな行を挿入し、データを入力する）
    # para = doc.paragraphs[5]
    # para.insert_paragraph_before("申請者住所　伊達市〇〇町１２３－４")

    # doc.paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ここから下は表

    # tbl = doc.tables[0]
    # for row in tbl.rows:
    #     row_text = []
    #     for cell in row.cells:
    #         row_text.append(cell.text)
    #     # print(row_text)
    #     print("".join(row_text).replace("\u3000", "_"))

    # 計算式@マークの部分はDBから受け取る
    for productinform in form.product_array:
        product = productinform.get_product()
        pass
        if product is None:
            continue
        yen_per_kg: int = product.yen_per_kg
        kg_per_1a = product.kg_per_1a
        cost = 0
        hectal = productinform.amount
        # 生産量 = @経営規模(広さ) * @単位規模当たりの生産量
        pro_amount = hectal * kg_per_1a
        # 農業粗収益 = 生産量 * @単価
        gross_profit = pro_amount * yen_per_kg
        # 農業所得 = 農業粗収益 - @農業経営費
        income = pro_amount - cost

    new_text_list1 = ["ジャガ", "10a", "1000kg"]
    new_text_list2 = ["人参", "8a", "800kg"]
    new_text_list3 = ["", "", ""]
    new_text_list4 = ["", "", ""]
    new_text_list5 = ["", "", ""]
    new_text_list6 = ["", "", ""]
    new_text_list7 = ["", "", ""]
    new_text_list8 = ["", "", ""]

    replace_words = ["作物", "面積", "生産量"]
    replace_lists = []
    for i in range(8):
        replace_list = []
        for word in replace_words:
            replace_list.append(f"{word}{i+1}")
        replace_lists.append(replace_list)

    replace_List1 = ["作物１", "面積１", "生産量１"]
    replace_List2 = ["作物２", "面積２", "生産量２"]
    replace_List3 = ["作物３", "面積３", "生産量３"]
    replace_List4 = ["作物４", "面積４", "生産量４"]
    replace_List5 = ["作物５", "面積５", "生産量５"]
    replace_List6 = ["作物６", "面積６", "生産量６"]
    replace_List7 = ["作物７", "面積７", "生産量７"]
    replace_List8 = ["作物８", "面積８", "生産量８"]

    for i in range(len(replace_List1)):
        table_replace(replace_List1[i], str(new_text_list1[i]))

    for i in range(len(replace_List2)):
        table_replace(replace_List2[i], str(new_text_list2[i]))

    for i in range(len(replace_List3)):
        table_replace(replace_List3[i], str(new_text_list3[i]))

    for i in range(len(replace_List4)):
        table_replace(replace_List4[i], str(new_text_list4[i]))

    for i in range(len(replace_List5)):
        table_replace(replace_List5[i], str(new_text_list5[i]))

    for i in range(len(replace_List6)):
        table_replace(replace_List6[i], str(new_text_list6[i]))

    for i in range(len(replace_List7)):
        table_replace(replace_List7[i], str(new_text_list7[i]))

    for i in range(len(replace_List8)):
        table_replace(replace_List8[i], str(new_text_list8[i]))

        # 年間農業粗収益を入力
        table_replace("@income", str(income))

    # 農業粗収益
    table_replace("農業粗収益の計算", str())
    makedirs("./tmp/", exist_ok=True)

    # 別名保存
    doc.save("./tmp/Result.docx")
    print("./tmp/Result.docx に保存")

    # 保存後のファイル中身確認
    doc = open_docx("./tmp/Result.docx")  # ファイル名とパスを正確に指定


def replace(len: int, old_text: str, new_text: str):
    rep = DOC.paragraphs[len]
    t = rep.text
    t = t.replace(old_text, new_text)
    rep.text = t
    return


# 表の中の文字を置き換える
def table_replace(old_text: str, new_text: str):
    for table in DOC.tables:  # 文書内のすべての表をループ
        for row in table.rows:  # 各表のすべての行をループ
            for cell in row.cells:  # 各行のすべてのセルをループ
                # 結合されたセルの最初のセルでのみテキストを置き換え
                if (
                    cell.text and old_text in cell.text
                ):  # テキストが空でない場合に置き換え
                    cell.text = cell.text.replace(old_text, new_text)
    return


if __name__ == "__main__":
    main(Form(id=5, product_array=[Form.ProductInForm(id=4, amount=3)], manpower=1))
